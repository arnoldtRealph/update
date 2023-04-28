import streamlit as st
import os
import io
import docx
from datetime import date
import streamlit.components.v1 as components
from streamlit_extras.colored_header import colored_header
from streamlit_extras.app_logo import add_logo
import streamlit_analytics

streamlit_analytics.start_tracking()
# Set page title and icon
st.set_page_config(page_title="Lesson Plan Creator", page_icon=":books:")


# add app logo
add_logo("IMAGES/wapen.png", height=150)

st.title("SAUL DAMON HIGH SCHOOL")

# using streamlit extras colored eader
colored_header(
    label="LESSON PLANNER",
    description="This interactive tool helps you create fast and easy lesson plans.",
    color_name="green-70",
)

# Create input fields
st.write("Please fill in the following fields to create a new lesson plan:")
st.write("")
subject = st.text_input("SUBJECT")
lesson_title = st.text_input("LESSON TITLE")
grade = st.selectbox("GRADE", ["9", "10", "11", "12"])
st.subheader("CHOOSE YOUR START AND END DATE")
start_date = st.date_input("FROM", value=date.today())
end_date = st.date_input("TO", value=date.today())
lesson_objective = st.text_area("LESSON OBJECTIVES")
lesson_introduction = st.text_area("INTRODUCTION")
lesson_activities = st.text_area("LESSON ACTIVITIES")
materials_needed = st.text_area("MATERIAL NEEDED")
homework = st.text_area("HOMEWORK ACTIVITIES")
notes = st.text_area("NOTES")
st.subheader("TEACHER INFORMATION")
teacher_name = st.text_input("INITIALS")
teacher_surname = st.text_input("SURNAME")
st.write("")
st.write("Created by Mr. A.R Visagie @ Saul Damon High School")

# Create save button
if st.button("Create Lesson Plan"):
    # Create a new Word document
    document = docx.Document()
    # Add input values to the document
    document.add_heading(subject.upper(), level=0)
    document.add_paragraph("")
    document.add_paragraph("LESSON TITLE: " + lesson_title)
    document.add_paragraph("GRADE: " + grade)
    document.add_paragraph("FROM: " + str(start_date))
    document.add_paragraph("TO: " + str(end_date))
    document.add_paragraph("")
    document.add_heading("LESSON OBJECTIVES", level=1)
    document.add_paragraph(lesson_objective)
    document.add_heading("INTRODUCTION", level=1)
    document.add_paragraph(lesson_introduction)
    document.add_heading("LESSON ACTIVITIES", level=1)
    document.add_paragraph(lesson_activities)
    document.add_heading("MATERIAL NEEDED", level=1)
    document.add_paragraph(materials_needed)
    document.add_heading("HOMEWORK ACTIVITIES", level=1)
    document.add_paragraph(homework)
    document.add_heading("NOTES", level=1)
    document.add_paragraph(notes)
    document.add_paragraph("")
    document.add_paragraph("INITIALS: " + teacher_name)
    document.add_paragraph("SURNAME: " + teacher_surname)
    document.add_paragraph("SIGNATURE: ")

    # Save document to BytesIO object
    with io.BytesIO() as output:
        document.save(output)
        output.seek(0)
        # Create download button
        st.download_button(
            label="Download Lesson Plan",
            data=output,
            file_name="lesson_plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("Your lesson plan has been created. Click the download button to save the file.")

     # Add a fun element
    st.balloons()

streamlit_analytics.stop_tracking()
