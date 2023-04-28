import streamlit as st
import os
import io
import docx
from datetime import date
import streamlit.components.v1 as components
from PIL import Image
from streamlit_extras.colored_header import colored_header
from streamlit_extras.app_logo import add_logo
import streamlit_analytics
import requests
from datetime import datetime

streamlit_analytics.start_tracking()


Header_image = Image.open("IMAGES/header.png")

# Set page title and icon
st.set_page_config(page_title="Lesson Plan Creator", page_icon=":books:", layout= "wide")
st.image("IMAGES/header.png")




# add app logo
add_logo("IMAGES/wapen.png", height=150)


# using streamlit extras colored eader

colored_header(
    label="LESBEPLANNER",
    description="Hierdie interaktiewe webblad help u om maklik en vinnig lesplanne te skep.",
    color_name="red-70",
)


# Create input fields

st.write("Vul asseblief die volgende velde in om 'n nuwe lesplan te skep:")
st.write("")
subject = st.text_input("VAK")
lesson_title = st.text_input("LES TITEL")
grade = st.selectbox("GRAAD", ["9", "10", "11", "12"])
st.subheader("KIES JOU BEGIN- EN EINDDATUM")
start_date = st.date_input("VANAF", value=date.today())
end_date = st.date_input("TOT", value=date.today())
lesson_objective = st.text_area("LES DOELWITTE")
lesson_introduction = st.text_area("INLEIDING")
lesson_activities = st.text_area("LES AKTIWITEITE")
materials_needed = st.text_area("MATERIAAL BENODIG")
homework = st.text_area("HUISWERK")
notes = st.text_area("NOTAS")
st.subheader("OPVOEDER BESONDERHEDE")
teacher_name = st.text_input("VOORLETTERS")
teacher_surname = st.text_input("VAN")
st.write("")
st.write("Created by Mr. A.R Visagie @ Saul Damon High School")

# Create save button
if st.button("Create Lesson Plan"):
    # Create a new Word document
    document = docx.Document()
    # Add input values to the document
    document.add_heading(subject.upper(), level=0)
    document.add_paragraph("")
    document.add_paragraph("LES TITEL: " + lesson_title)
    document.add_paragraph("GRAAD: " + grade)
    document.add_paragraph("VANAF: " + str(start_date))
    document.add_paragraph("TOT: " + str(end_date))
    document.add_paragraph("")
    document.add_heading("LES DOELWITTE", level=1)
    document.add_paragraph(lesson_objective)
    document.add_heading("INLEIDING", level=1)
    document.add_paragraph(lesson_introduction)
    document.add_heading("LES AKTIWITEITE", level=1)
    document.add_paragraph(lesson_activities)
    document.add_heading("MATERIAAL BENODIG", level=1)
    document.add_paragraph(materials_needed)
    document.add_heading("HUISWERK", level=1)
    document.add_paragraph(homework)
    document.add_heading("NOTAS", level=1)
    document.add_paragraph(notes)
    document.add_paragraph("")
    document.add_paragraph("VOORLETTERS: " + teacher_name)
    document.add_paragraph("VAN: " + teacher_surname)
    document.add_paragraph("HANDTEKENNG: ")

    # Save document to BytesIO object
    with io.BytesIO() as output:
        document.save(output)
        output.seek(0)
        # Create download button
        st.download_button(
            label="Download Lesson Plan",
            data=output,
            file_name="lesson_plan.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("Your lesson plan has been created. Click the download button to save the file.")


     # Add a fun element
    st.balloons()

streamlit_analytics.stop_tracking()

#Temperature widget

latitude = st.secrets.my_secrets.latitude
longitude = st.secrets.my_secrets.longitude
API_KEY = st.secrets.my_secrets.API_KEY

url = f'https://api.openweathermap.org/data/2.5/weather?lat={latitude}&lon={longitude}&appid={API_KEY}'

response = requests.get(url)

if response.status_code == 200:
    data = response.json()
    temperature = data['main']['temp'] - 273.15  # Convert from Kelvin to Celsius
    description = data['weather'][0]['description']
    icon_code = data['weather'][0]['icon']
    icon_url = f"http://openweathermap.org/img/w/{icon_code}.png"

    st.title("Upington Weather")
    st.write(f"Last updated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    st.image(icon_url, width=100)
    st.markdown(f'<p style="color: #F2C94C;">{temperature:.1f}°C</p>', unsafe_allow_html=True)
    st.markdown(f'<p style="color: #F2C94C;">{description.capitalize()}</p>', unsafe_allow_html=True)
    st.markdown('<style>div.row-widget.stRadio > div{flex-direction:row;}</style>', unsafe_allow_html=True)
else:
    st.write('Error retrieving temperature data.')







